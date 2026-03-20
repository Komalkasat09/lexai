"""
Document Extraction Module
Handles extraction of text from PDF and DOCX files.
Only supports typed documents - rejects scanned/handwritten content.
"""

import fitz  # PyMuPDF
from docx import Document
from typing import Tuple, Optional
import re


class DocumentExtractor:
    """Extracts text from PDF and DOCX files for contract analysis."""
    
    # Minimum text density to consider document as typed (not scanned)
    MIN_TEXT_DENSITY = 50  # characters per page
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Tuple[str, dict]:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            ValueError: If document appears to be scanned
        """
        try:
            doc = fitz.open(file_path)
            
            # Extract metadata
            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "format": "PDF"
            }
            
            # Extract text from all pages
            full_text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                full_text += page_text + "\n\n"
            
            doc.close()
            
            # Validate that document is typed (not scanned)
            avg_chars_per_page = len(full_text.strip()) / max(metadata["page_count"], 1)
            if avg_chars_per_page < DocumentExtractor.MIN_TEXT_DENSITY:
                raise ValueError(
                    "Document appears to be scanned or contains insufficient text. "
                    "Only typed PDF documents are supported."
                )
            
            return full_text.strip(), metadata
            
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Tuple[str, dict]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            ValueError: If extraction fails
        """
        try:
            doc = Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "page_count": None,  # DOCX doesn't reliably track page count
                "title": core_props.title or "",
                "author": core_props.author or "",
                "format": "DOCX"
            }
            
            # Extract text from all paragraphs
            full_text = ""
            for paragraph in doc.paragraphs:
                # Preserve paragraph structure
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
            
            # Also extract text from tables (contracts often have tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text += row_text + "\n"
            
            # Validate that document has content
            if len(full_text.strip()) < DocumentExtractor.MIN_TEXT_DENSITY:
                raise ValueError(
                    "Document appears to be empty or contains insufficient text."
                )
            
            return full_text.strip(), metadata
            
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract(file_path: str, file_extension: str) -> Tuple[str, dict]:
        """
        Main extraction method - routes to appropriate extractor.
        
        Args:
            file_path: Path to document file
            file_extension: File extension (.pdf or .docx)
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            ValueError: If file type not supported or extraction fails
        """
        file_extension = file_extension.lower().strip('.')
        
        if file_extension == 'pdf':
            return DocumentExtractor.extract_from_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            if file_extension == 'doc':
                raise ValueError(
                    "Legacy .doc format is not supported. "
                    "Please convert to .docx format."
                )
            return DocumentExtractor.extract_from_docx(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: .{file_extension}. "
                "Only PDF and DOCX files are supported."
            )
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and normalizing.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove page numbers (common pattern: "Page 1 of 10")
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n')
        
        return text.strip()
